"""Tests for ergodix.importers.docx — Word/Scrivener .docx → Pandoc-Markdown.

Per ADR 0015 §2: `.docx` is a first-class import target, extracted via
`python-docx` (already a runtime dep). Coverage in v1 mirrors the gdocs
extractor: paragraphs, headings (Heading 1-6 + Title), bold/italic,
inline links, bulleted lists. Tables / images / footnotes deferred to
chunks 6+ per ADR 0015's chunk plan.

Tests build sample `.docx` files in `tmp_path` using `python-docx`
itself, then call `docx.extract` and assert on the rendered Markdown.
No network, no Drive, no OAuth.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from ergodix.importers import docx as docx_importer


def _build_docx(tmp_path: Path, build_fn: object) -> Path:
    """Build a .docx via python-docx, return the path."""
    from docx import Document

    document = Document()
    build_fn(document)
    out = tmp_path / "test.docx"
    document.save(str(out))
    return out


# ─── module shape ──────────────────────────────────────────────────────────


def test_module_declares_name_and_extensions() -> None:
    assert docx_importer.NAME == "docx"
    assert ".docx" in docx_importer.EXTENSIONS


def test_module_registered_in_importer_registry() -> None:
    from ergodix import importers

    assert "docx" in importers.available_importers()
    assert importers.extension_to_importer(".docx") is not None
    assert importers.extension_to_importer(".docx").name == "docx"


# ─── extract — happy paths ─────────────────────────────────────────────────


def test_renders_normal_paragraph(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("Hello world.")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "Hello world." in md


def test_renders_h1_through_h6(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        for level in range(1, 7):
            doc.add_heading(f"Heading {level}", level=level)

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "# Heading 1" in md
    assert "## Heading 2" in md
    assert "### Heading 3" in md
    assert "#### Heading 4" in md
    assert "##### Heading 5" in md
    assert "###### Heading 6" in md


def test_title_renders_as_h1(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        # python-docx's add_heading(level=0) maps to "Title" style.
        doc.add_heading("My Title", level=0)

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "# My Title" in md


def test_bold_run(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        para = doc.add_paragraph("It was ")
        run = para.add_run("bold")
        run.bold = True
        para.add_run(".")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "It was **bold**." in md


def test_italic_run(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        para = doc.add_paragraph("Walking ")
        run = para.add_run("alone")
        run.italic = True
        para.add_run(" home.")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "Walking *alone* home." in md


def test_bold_and_italic_combined(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        para = doc.add_paragraph()
        run = para.add_run("loud")
        run.bold = True
        run.italic = True

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "***loud***" in md


def test_paragraphs_separated_by_blank_line(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "First paragraph.\n\nSecond paragraph." in md


def test_bulleted_list_items(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("first", style="List Bullet")
        doc.add_paragraph("second", style="List Bullet")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "- first" in md
    assert "- second" in md


def test_empty_document_returns_empty_string(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        pass  # No paragraphs added.

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    # python-docx adds one empty paragraph automatically — should still render empty.
    assert md == ""


def test_skips_empty_paragraphs(tmp_path: Path) -> None:
    """Empty paragraphs (just whitespace / no runs) shouldn't emit blank lines."""

    def build(doc: object) -> None:
        doc.add_paragraph("Real content.")
        doc.add_paragraph("")  # empty
        doc.add_paragraph("More content.")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "Real content." in md
    assert "More content." in md
    # No triple-blank-line gap.
    assert "\n\n\n" not in md


def test_extract_raises_on_missing_file(tmp_path: Path) -> None:
    missing = tmp_path / "absent.docx"
    with pytest.raises(FileNotFoundError):
        docx_importer.extract(missing)


def test_extract_accepts_unknown_kwargs(tmp_path: Path) -> None:
    """The orchestrator passes docs_service to every importer; non-gdocs
    importers must accept (and ignore) it without raising."""

    def build(doc: object) -> None:
        doc.add_paragraph("Hello.")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path, docs_service=object())
    assert "Hello." in md


def test_unknown_paragraph_styles_pass_through_as_plain_text(tmp_path: Path) -> None:
    """A paragraph with an unrecognized style (custom Word template) renders
    as plain markdown rather than dropping content."""

    def build(doc: object) -> None:
        # Use a built-in but unhandled-by-our-renderer style.
        doc.add_paragraph("Normal-style content.")
        doc.add_paragraph("Subtitle content.", style="Subtitle")

    path = _build_docx(tmp_path, build)
    md = docx_importer.extract(path)
    assert "Normal-style content." in md
    assert "Subtitle content." in md


# ─── Embedded images (chunk 6) ─────────────────────────────────────────────


# Minimal valid 1x1 black PNG — pre-encoded so tests don't need PIL/Pillow.
_TINY_PNG = bytes.fromhex(
    "89504e470d0a1a0a"
    "0000000d49484452"
    "0000000100000001"
    "0100000000376ef9"
    "240000000a494441"
    "54789c6300000000"
    "0200015fc7c2bd00"
    "00000049454e44ae"
    "426082"
)


def _make_docx_with_image(tmp_path: Path) -> Path:
    """Build a .docx with one normal paragraph + one embedded image."""
    from docx import Document
    from docx.shared import Inches

    image_source = tmp_path / "tiny.png"
    image_source.write_bytes(_TINY_PNG)

    document = Document()
    document.add_paragraph("Before the image.")
    document.add_picture(str(image_source), width=Inches(1.0))
    document.add_paragraph("After the image.")

    out = tmp_path / "with-image.docx"
    document.save(str(out))
    return out


def test_extract_writes_image_bytes_to_media_dir(tmp_path: Path) -> None:
    """When `media_dir` is provided, images embedded in the .docx are
    saved as files there. Filename follows `img-NNN.<ext>` per ADR 0015 §3."""
    docx_path = _make_docx_with_image(tmp_path)
    media = tmp_path / "_media" / "with-image"

    md = docx_importer.extract(docx_path, media_dir=media)

    images = sorted(media.iterdir()) if media.exists() else []
    assert len(images) == 1
    assert images[0].suffix == ".png"
    assert images[0].read_bytes() == _TINY_PNG
    # Markdown contains a reference to the saved image.
    assert "![" in md
    assert images[0].name in md


def test_extract_skips_image_extraction_when_media_dir_omitted(tmp_path: Path) -> None:
    """Without `media_dir`, the importer doesn't fail on embedded images;
    it just doesn't write them. Mirrors the behavior the orchestrator
    relies on for `--check` (dry-run)."""
    docx_path = _make_docx_with_image(tmp_path)

    # Should not raise, even though there's an image.
    md = docx_importer.extract(docx_path)

    assert "Before the image." in md
    assert "After the image." in md
    assert not (tmp_path / "_media").exists()


def test_extract_handles_multiple_images_in_one_doc(tmp_path: Path) -> None:
    """Two distinct embedded images get sequential filenames
    (img-001, img-002)."""
    from docx import Document
    from docx.shared import Inches

    # Two distinct image sources so python-docx doesn't dedupe via a
    # single shared relationship. The second is the same byte content
    # plus a tiny tail — different content fingerprint, different part.
    image_a = tmp_path / "tiny-a.png"
    image_b = tmp_path / "tiny-b.png"
    image_a.write_bytes(_TINY_PNG)
    image_b.write_bytes(_TINY_PNG[:-3] + b"\x00\x00\x00" + _TINY_PNG[-3:])

    document = Document()
    document.add_paragraph("Para 1.")
    document.add_picture(str(image_a), width=Inches(1.0))
    document.add_paragraph("Para 2.")
    document.add_picture(str(image_b), width=Inches(1.0))
    document.add_paragraph("Para 3.")

    docx_path = tmp_path / "two-images.docx"
    document.save(str(docx_path))

    media = tmp_path / "_media" / "two-images"
    docx_importer.extract(docx_path, media_dir=media)

    images = sorted(media.iterdir())
    assert len(images) == 2
    assert images[0].name.startswith("img-001")
    assert images[1].name.startswith("img-002")
