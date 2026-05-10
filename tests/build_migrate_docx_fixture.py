"""
One-shot builder for the migrate-fixture .docx file(s).

Regenerate the checked-in `.docx` fixture(s) under
`examples/migrate-fixture/` when their expected content changes. The
fixtures are stable binaries (committed to git) so the hermetic e2e
tests at `tests/test_migrate_fixture.py` exercise the full migrate run
against real Word/Scrivener-shaped input without needing Drive or a
network round-trip.

Lives under `tests/` (not inside the fixture itself) so the migrate
walker doesn't see the `.py` builder as an out-of-scope corpus file.

Usage (from repo root):

    python tests/build_migrate_docx_fixture.py

Re-run after touching this file or whenever python-docx's output shape
changes. The resulting `.docx` is checked in alongside the rest of the
fixture.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE_DIR = Path(__file__).parent.parent / "examples" / "migrate-fixture"


def build_chapter_2() -> None:
    """Chapter 2: a `.docx` with paragraphs, headings, bold/italic, lists."""
    document = Document()

    document.add_heading("Chapter 2 — The Mirrored Hall", level=1)

    document.add_paragraph(
        "Theron stepped onto the floor of the great hall and the floor "
        "stepped back, mirror-bright, mirror-quiet."
    )

    para = document.add_paragraph("In the centre stood a man who was ")
    run = para.add_run("not a man")
    run.italic = True
    para.add_run(" — and who knew, ")
    run = para.add_run("knew")
    run.bold = True
    para.add_run(", that Theron was watching.")

    document.add_heading("What Theron carried", level=2)

    document.add_paragraph("a length of rope", style="List Bullet")
    document.add_paragraph("a folded letter, sealed in green wax", style="List Bullet")
    document.add_paragraph("the certainty that one of these would matter", style="List Bullet")

    document.add_paragraph("He would learn which one only after the choice was made.")

    target = FIXTURE_DIR / "Book 1" / "Chapter 2.docx"
    document.save(str(target))
    print(f"wrote {target}")


def main() -> None:
    build_chapter_2()


if __name__ == "__main__":
    main()
